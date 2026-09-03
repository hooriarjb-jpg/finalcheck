import streamlit as st
import re, json
from io import BytesIO
from difflib import SequenceMatcher

st.set_page_config(page_title="FinalCheck V2", page_icon="✓", layout="wide")

st.markdown("""
<style>
.block-container {max-width: 1120px; padding-top: 2rem;}
.pass {padding:14px;border-radius:10px;background:#eaf7ee;margin:8px 0}
.conflict {padding:14px;border-radius:10px;background:#fdecec;margin:8px 0}
.review {padding:14px;border-radius:10px;background:#fff6df;margin:8px 0}
</style>
""", unsafe_allow_html=True)

st.title("FinalCheck")
st.caption("Pre-flight fact checking for communications documents — V2")

def extract_text(uploaded):
    if uploaded is None: return ""
    data = uploaded.getvalue()
    name = uploaded.name.lower()
    if name.endswith((".txt",".md")):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        from docx import Document
        return "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        return "\n".join((p.extract_text() or "") for p in PdfReader(BytesIO(data)).pages)
    return ""

MONTHS = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"

def unique(xs):
    seen=[]; keys=set()
    for x in xs:
        x=re.sub(r"\s+"," ",x.strip().rstrip(".,;:"))
        k=x.lower()
        if x and k not in keys:
            keys.add(k); seen.append(x)
    return seen

def extract(text):
    return {
      "Date": unique(re.findall(rf"\b{MONTHS}\s+\d{{1,2}}(?:,\s*\d{{4}})?\b|\b\d{{4}}[-/]\d{{1,2}}[-/]\d{{1,2}}\b", text, re.I)),
      "Time": unique(re.findall(r"\b\d{1,2}(?::\d{2})?\s*(?:a\.?m\.?|p\.?m\.?)\b", text, re.I)),
      "Price": unique(re.findall(r"(?:CAD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?", text, re.I)),
      "URL": unique(re.findall(r"https?://[^\s<>\]\)]+|www\.[^\s<>\]\)]+", text, re.I)),
      "Email": unique(re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)),
    }

def canon(category, value):
    s=value.lower().strip()
    s=s.replace("a.m.","am").replace("p.m.","pm").replace("a.m","am").replace("p.m","pm")
    s=re.sub(r"\s+"," ",s)
    if category=="Time":
        m=re.search(r"(\d{1,2})(?::(\d{2}))?\s*(am|pm)",s)
        if m:
            h=int(m.group(1)); minute=int(m.group(2) or 0); ap=m.group(3)
            if ap=="pm" and h!=12: h+=12
            if ap=="am" and h==12: h=0
            return f"{h:02}:{minute:02}"
    if category=="Price":
        return re.sub(r"[^\d.]","",s)
    if category in ("URL","Email"):
        return s.rstrip("/")
    return re.sub(r"[^a-z0-9]","",s)

def pair_values(src, drf, category):
    results=[]
    remaining=list(drf)

    # First remove equivalent values as PASS.
    unmatched=[]
    for s in src:
        exact=next((d for d in remaining if canon(category,d)==canon(category,s)), None)
        if exact is not None:
            results.append(("PASS",category,s,exact,"Matches approved source."))
            remaining.remove(exact)
        else:
            unmatched.append(s)

    # If source and draft both have a value of a high-risk category, a changed
    # value is a conflict rather than a generic omission.
    for s in unmatched:
        if remaining and category in ("Date","Time","Price"):
            candidate=remaining.pop(0)
            results.append(("CONFLICT",category,s,candidate,"Approved and final values conflict."))
        else:
            results.append(("REVIEW",category,s,None,"Present in approved source but not found in final draft."))

    # Extra high-risk values appearing only in the draft deserve review.
    for d in remaining:
        if category in ("Date","Time","Price"):
            results.append(("REVIEW",category,"Not in approved source",d,"Final draft contains an additional value."))
    return results

def compare(source,draft,required):
    S,D=extract(source),extract(draft)
    rows=[]
    for cat in ["Date","Time","Price","URL","Email"]:
        rows += pair_values(S[cat],D[cat],cat)
    for phrase in required:
        if phrase.lower() in draft.lower():
            rows.append(("PASS","Required wording",phrase,phrase,"Required wording is present."))
        else:
            rows.append(("CONFLICT","Required wording",phrase,"Missing","Required wording is missing."))
    return rows

c1,c2=st.columns(2)
with c1:
    st.subheader("1. Approved source")
    sf=st.file_uploader("Upload approved source",type=["pdf","docx","txt","md"])
    sp=st.text_area("Or paste approved source",height=180)
with c2:
    st.subheader("2. Final draft")
    df=st.file_uploader("Upload final draft",type=["pdf","docx","txt","md"])
    dp=st.text_area("Or paste final draft",height=180)

st.subheader("3. Required wording (optional)")
req=st.text_area("Only add wording that must appear in the final version — one item per line",height=90)

if st.button("Run FinalCheck",type="primary",use_container_width=True):
    source=sp.strip() or extract_text(sf)
    draft=dp.strip() or extract_text(df)
    if not source or not draft:
        st.error("Please provide both documents.")
    else:
        required_items=[x.strip() for x in req.splitlines() if x.strip()]
        source_entities=extract(source)
        draft_entities=extract(draft)
        rows=compare(source,draft,required_items)

        with st.expander("Extraction check — what FinalCheck read from each document"):
            left,right=st.columns(2)
            with left:
                st.markdown("**Approved source**")
                for cat, vals in source_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            with right:
                st.markdown("**Final draft**")
                for cat, vals in draft_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            st.caption("If a fact is missing here, the extraction step—not the comparison step—is the problem.")

        conflicts=[r for r in rows if r[0]=="CONFLICT"]
        reviews=[r for r in rows if r[0]=="REVIEW"]
        passes=[r for r in rows if r[0]=="PASS"]

        a,b,c=st.columns(3)
        a.metric("Conflicts",len(conflicts))
        b.metric("Needs review",len(reviews))
        c.metric("Passed",len(passes))

        if conflicts: st.error("Conflicting information found.")
        elif reviews: st.warning("No direct conflicts found, but some source information needs review.")
        else: st.success("No factual conflicts detected.")

        for status,cat,s,d,msg in conflicts+reviews+passes:
            icon={"CONFLICT":"🔴","REVIEW":"🟡","PASS":"🟢"}[status]
            with st.expander(f"{icon} {status} · {cat} · {s}", expanded=status=="CONFLICT"):
                st.write(f"**Approved source:** {s}")
                st.write(f"**Final draft:** {d or 'Not found'}")
                st.write(msg)

        report=[{"status":x[0],"category":x[1],"approved":x[2],"draft":x[3],"reason":x[4]} for x in rows]
        st.download_button("Download QA report",json.dumps(report,indent=2,ensure_ascii=False),
                           "finalcheck_v2_report.json","application/json",use_container_width=True)

st.divider()
st.caption("V2 separates direct conflicts from omissions/review items and recognizes equivalent time formatting.")
import streamlit as st
import re, json
from io import BytesIO
from difflib import SequenceMatcher

st.set_page_config(page_title="FinalCheck V2", page_icon="✓", layout="wide")

st.markdown("""
<style>
.block-container {max-width: 1120px; padding-top: 2rem;}
.pass {padding:14px;border-radius:10px;background:#eaf7ee;margin:8px 0}
.conflict {padding:14px;border-radius:10px;background:#fdecec;margin:8px 0}
.review {padding:14px;border-radius:10px;background:#fff6df;margin:8px 0}
</style>
""", unsafe_allow_html=True)

st.title("FinalCheck")
st.caption("Pre-flight fact checking for communications documents — V2")

def extract_text(uploaded):
    if uploaded is None: return ""
    data = uploaded.getvalue()
    name = uploaded.name.lower()
    if name.endswith((".txt",".md")):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        from docx import Document
        return "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        return "\n".join((p.extract_text() or "") for p in PdfReader(BytesIO(data)).pages)
    return ""

MONTHS = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"

def unique(xs):
    seen=[]; keys=set()
    for x in xs:
        x=re.sub(r"\s+"," ",x.strip().rstrip(".,;:"))
        k=x.lower()
        if x and k not in keys:
            keys.add(k); seen.append(x)
    return seen

def extract(text):
    return {
      "Date": unique(re.findall(rf"\b{MONTHS}\s+\d{{1,2}}(?:,\s*\d{{4}})?\b|\b\d{{4}}[-/]\d{{1,2}}[-/]\d{{1,2}}\b", text, re.I)),
      "Time": unique(re.findall(r"\b\d{1,2}(?::\d{2})?\s*(?:a\.?m\.?|p\.?m\.?)\b", text, re.I)),
      "Price": unique(re.findall(r"(?:CAD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?", text, re.I)),
      "URL": unique(re.findall(r"https?://[^\s<>\]\)]+|www\.[^\s<>\]\)]+", text, re.I)),
      "Email": unique(re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)),
    }

def canon(category, value):
    s=value.lower().strip()
    s=s.replace("a.m.","am").replace("p.m.","pm").replace("a.m","am").replace("p.m","pm")
    s=re.sub(r"\s+"," ",s)
    if category=="Time":
        m=re.search(r"(\d{1,2})(?::(\d{2}))?\s*(am|pm)",s)
        if m:
            h=int(m.group(1)); minute=int(m.group(2) or 0); ap=m.group(3)
            if ap=="pm" and h!=12: h+=12
            if ap=="am" and h==12: h=0
            return f"{h:02}:{minute:02}"
    if category=="Price":
        return re.sub(r"[^\d.]","",s)
    if category in ("URL","Email"):
        return s.rstrip("/")
    return re.sub(r"[^a-z0-9]","",s)

def pair_values(src, drf, category):
    results=[]
    remaining=list(drf)

    # First remove equivalent values as PASS.
    unmatched=[]
    for s in src:
        exact=next((d for d in remaining if canon(category,d)==canon(category,s)), None)
        if exact is not None:
            results.append(("PASS",category,s,exact,"Matches approved source."))
            remaining.remove(exact)
        else:
            unmatched.append(s)

    # If source and draft both have a value of a high-risk category, a changed
    # value is a conflict rather than a generic omission.
    for s in unmatched:
        if remaining and category in ("Date","Time","Price"):
            candidate=remaining.pop(0)
            results.append(("CONFLICT",category,s,candidate,"Approved and final values conflict."))
        else:
            results.append(("REVIEW",category,s,None,"Present in approved source but not found in final draft."))

    # Extra high-risk values appearing only in the draft deserve review.
    for d in remaining:
        if category in ("Date","Time","Price"):
            results.append(("REVIEW",category,"Not in approved source",d,"Final draft contains an additional value."))
    return results

def compare(source,draft,required):
    S,D=extract(source),extract(draft)
    rows=[]
    for cat in ["Date","Time","Price","URL","Email"]:
        rows += pair_values(S[cat],D[cat],cat)
    for phrase in required:
        if phrase.lower() in draft.lower():
            rows.append(("PASS","Required wording",phrase,phrase,"Required wording is present."))
        else:
            rows.append(("CONFLICT","Required wording",phrase,"Missing","Required wording is missing."))
    return rows

c1,c2=st.columns(2)
with c1:
    st.subheader("1. Approved source")
    sf=st.file_uploader("Upload approved source",type=["pdf","docx","txt","md"])
    sp=st.text_area("Or paste approved source",height=180)
with c2:
    st.subheader("2. Final draft")
    df=st.file_uploader("Upload final draft",type=["pdf","docx","txt","md"])
    dp=st.text_area("Or paste final draft",height=180)

st.subheader("3. Required wording (optional)")
req=st.text_area("Only add wording that must appear in the final version — one item per line",height=90)

if st.button("Run FinalCheck",type="primary",use_container_width=True):
    source=sp.strip() or extract_text(sf)
    draft=dp.strip() or extract_text(df)
    if not source or not draft:
        st.error("Please provide both documents.")
    else:
        required_items=[x.strip() for x in req.splitlines() if x.strip()]
        source_entities=extract(source)
        draft_entities=extract(draft)
        rows=compare(source,draft,required_items)

        with st.expander("Extraction check — what FinalCheck read from each document"):
            left,right=st.columns(2)
            with left:
                st.markdown("**Approved source**")
                for cat, vals in source_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            with right:
                st.markdown("**Final draft**")
                for cat, vals in draft_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            st.caption("If a fact is missing here, the extraction step—not the comparison step—is the problem.")

        conflicts=[r for r in rows if r[0]=="CONFLICT"]
        reviews=[r for r in rows if r[0]=="REVIEW"]
        passes=[r for r in rows if r[0]=="PASS"]

        a,b,c=st.columns(3)
        a.metric("Conflicts",len(conflicts))
        b.metric("Needs review",len(reviews))
        c.metric("Passed",len(passes))

        if conflicts: st.error("Conflicting information found.")
        elif reviews: st.warning("No direct conflicts found, but some source information needs review.")
        else: st.success("No factual conflicts detected.")

        for status,cat,s,d,msg in conflicts+reviews+passes:
            icon={"CONFLICT":"🔴","REVIEW":"🟡","PASS":"🟢"}[status]
            with st.expander(f"{icon} {status} · {cat} · {s}", expanded=status=="CONFLICT"):
                st.write(f"**Approved source:** {s}")
                st.write(f"**Final draft:** {d or 'Not found'}")
                st.write(msg)

        report=[{"status":x[0],"category":x[1],"approved":x[2],"draft":x[3],"reason":x[4]} for x in rows]
        st.download_button("Download QA report",json.dumps(report,indent=2,ensure_ascii=False),
                           "finalcheck_v2_report.json","application/json",use_container_width=True)

st.divider()
st.caption("V2 separates direct conflicts from omissions/review items and recognizes equivalent time formatting.")
import streamlit as st
import re, json
from io import BytesIO
from difflib import SequenceMatcher

st.set_page_config(page_title="FinalCheck V2", page_icon="✓", layout="wide")

st.markdown("""
<style>
.block-container {max-width: 1120px; padding-top: 2rem;}
.pass {padding:14px;border-radius:10px;background:#eaf7ee;margin:8px 0}
.conflict {padding:14px;border-radius:10px;background:#fdecec;margin:8px 0}
.review {padding:14px;border-radius:10px;background:#fff6df;margin:8px 0}
</style>
""", unsafe_allow_html=True)

st.title("FinalCheck")
st.caption("Pre-flight fact checking for communications documents — V2")

def extract_text(uploaded):
    if uploaded is None: return ""
    data = uploaded.getvalue()
    name = uploaded.name.lower()
    if name.endswith((".txt",".md")):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        from docx import Document
        return "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        return "\n".join((p.extract_text() or "") for p in PdfReader(BytesIO(data)).pages)
    return ""

MONTHS = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"

def unique(xs):
    seen=[]; keys=set()
    for x in xs:
        x=re.sub(r"\s+"," ",x.strip().rstrip(".,;:"))
        k=x.lower()
        if x and k not in keys:
            keys.add(k); seen.append(x)
    return seen

def extract(text):
    return {
      "Date": unique(re.findall(rf"\b{MONTHS}\s+\d{{1,2}}(?:,\s*\d{{4}})?\b|\b\d{{4}}[-/]\d{{1,2}}[-/]\d{{1,2}}\b", text, re.I)),
      "Time": unique(re.findall(r"\b\d{1,2}(?::\d{2})?\s*(?:a\.?m\.?|p\.?m\.?)\b", text, re.I)),
      "Price": unique(re.findall(r"(?:CAD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?", text, re.I)),
      "URL": unique(re.findall(r"https?://[^\s<>\]\)]+|www\.[^\s<>\]\)]+", text, re.I)),
      "Email": unique(re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)),
    }

def canon(category, value):
    s=value.lower().strip()
    s=s.replace("a.m.","am").replace("p.m.","pm").replace("a.m","am").replace("p.m","pm")
    s=re.sub(r"\s+"," ",s)
    if category=="Time":
        m=re.search(r"(\d{1,2})(?::(\d{2}))?\s*(am|pm)",s)
        if m:
            h=int(m.group(1)); minute=int(m.group(2) or 0); ap=m.group(3)
            if ap=="pm" and h!=12: h+=12
            if ap=="am" and h==12: h=0
            return f"{h:02}:{minute:02}"
    if category=="Price":
        return re.sub(r"[^\d.]","",s)
    if category in ("URL","Email"):
        return s.rstrip("/")
    return re.sub(r"[^a-z0-9]","",s)

def pair_values(src, drf, category):
    results=[]
    remaining=list(drf)

    # First remove equivalent values as PASS.
    unmatched=[]
    for s in src:
        exact=next((d for d in remaining if canon(category,d)==canon(category,s)), None)
        if exact is not None:
            results.append(("PASS",category,s,exact,"Matches approved source."))
            remaining.remove(exact)
        else:
            unmatched.append(s)

    # If source and draft both have a value of a high-risk category, a changed
    # value is a conflict rather than a generic omission.
    for s in unmatched:
        if remaining and category in ("Date","Time","Price"):
            candidate=remaining.pop(0)
            results.append(("CONFLICT",category,s,candidate,"Approved and final values conflict."))
        else:
            results.append(("REVIEW",category,s,None,"Present in approved source but not found in final draft."))

    # Extra high-risk values appearing only in the draft deserve review.
    for d in remaining:
        if category in ("Date","Time","Price"):
            results.append(("REVIEW",category,"Not in approved source",d,"Final draft contains an additional value."))
    return results

def compare(source,draft,required):
    S,D=extract(source),extract(draft)
    rows=[]
    for cat in ["Date","Time","Price","URL","Email"]:
        rows += pair_values(S[cat],D[cat],cat)
    for phrase in required:
        if phrase.lower() in draft.lower():
            rows.append(("PASS","Required wording",phrase,phrase,"Required wording is present."))
        else:
            rows.append(("CONFLICT","Required wording",phrase,"Missing","Required wording is missing."))
    return rows

c1,c2=st.columns(2)
with c1:
    st.subheader("1. Approved source")
    sf=st.file_uploader("Upload approved source",type=["pdf","docx","txt","md"])
    sp=st.text_area("Or paste approved source",height=180)
with c2:
    st.subheader("2. Final draft")
    df=st.file_uploader("Upload final draft",type=["pdf","docx","txt","md"])
    dp=st.text_area("Or paste final draft",height=180)

st.subheader("3. Required wording (optional)")
req=st.text_area("Only add wording that must appear in the final version — one item per line",height=90)

if st.button("Run FinalCheck",type="primary",use_container_width=True):
    source=sp.strip() or extract_text(sf)
    draft=dp.strip() or extract_text(df)
    if not source or not draft:
        st.error("Please provide both documents.")
    else:
        required_items=[x.strip() for x in req.splitlines() if x.strip()]
        source_entities=extract(source)
        draft_entities=extract(draft)
        rows=compare(source,draft,required_items)

        with st.expander("Extraction check — what FinalCheck read from each document"):
            left,right=st.columns(2)
            with left:
                st.markdown("**Approved source**")
                for cat, vals in source_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            with right:
                st.markdown("**Final draft**")
                for cat, vals in draft_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            st.caption("If a fact is missing here, the extraction step—not the comparison step—is the problem.")

        conflicts=[r for r in rows if r[0]=="CONFLICT"]
        reviews=[r for r in rows if r[0]=="REVIEW"]
        passes=[r for r in rows if r[0]=="PASS"]

        a,b,c=st.columns(3)
        a.metric("Conflicts",len(conflicts))
        b.metric("Needs review",len(reviews))
        c.metric("Passed",len(passes))

        if conflicts: st.error("Conflicting information found.")
        elif reviews: st.warning("No direct conflicts found, but some source information needs review.")
        else: st.success("No factual conflicts detected.")

        for status,cat,s,d,msg in conflicts+reviews+passes:
            icon={"CONFLICT":"🔴","REVIEW":"🟡","PASS":"🟢"}[status]
            with st.expander(f"{icon} {status} · {cat} · {s}", expanded=status=="CONFLICT"):
                st.write(f"**Approved source:** {s}")
                st.write(f"**Final draft:** {d or 'Not found'}")
                st.write(msg)

        report=[{"status":x[0],"category":x[1],"approved":x[2],"draft":x[3],"reason":x[4]} for x in rows]
        st.download_button("Download QA report",json.dumps(report,indent=2,ensure_ascii=False),
                           "finalcheck_v2_report.json","application/json",use_container_width=True)

st.divider()
st.caption("V2 separates direct conflicts from omissions/review items and recognizes equivalent time formatting.")
import streamlit as st
import re, json
from io import BytesIO
from difflib import SequenceMatcher

st.set_page_config(page_title="FinalCheck V2", page_icon="✓", layout="wide")

st.markdown("""
<style>
.block-container {max-width: 1120px; padding-top: 2rem;}
.pass {padding:14px;border-radius:10px;background:#eaf7ee;margin:8px 0}
.conflict {padding:14px;border-radius:10px;background:#fdecec;margin:8px 0}
.review {padding:14px;border-radius:10px;background:#fff6df;margin:8px 0}
</style>
""", unsafe_allow_html=True)

st.title("FinalCheck")
st.caption("Pre-flight fact checking for communications documents — V2")

def extract_text(uploaded):
    if uploaded is None: return ""
    data = uploaded.getvalue()
    name = uploaded.name.lower()
    if name.endswith((".txt",".md")):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        from docx import Document
        return "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        return "\n".join((p.extract_text() or "") for p in PdfReader(BytesIO(data)).pages)
    return ""

MONTHS = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"

def unique(xs):
    seen=[]; keys=set()
    for x in xs:
        x=re.sub(r"\s+"," ",x.strip().rstrip(".,;:"))
        k=x.lower()
        if x and k not in keys:
            keys.add(k); seen.append(x)
    return seen

def extract(text):
    return {
      "Date": unique(re.findall(rf"\b{MONTHS}\s+\d{{1,2}}(?:,\s*\d{{4}})?\b|\b\d{{4}}[-/]\d{{1,2}}[-/]\d{{1,2}}\b", text, re.I)),
      "Time": unique(re.findall(r"\b\d{1,2}(?::\d{2})?\s*(?:a\.?m\.?|p\.?m\.?)\b", text, re.I)),
      "Price": unique(re.findall(r"(?:CAD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?", text, re.I)),
      "URL": unique(re.findall(r"https?://[^\s<>\]\)]+|www\.[^\s<>\]\)]+", text, re.I)),
      "Email": unique(re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)),
    }

def canon(category, value):
    s=value.lower().strip()
    s=s.replace("a.m.","am").replace("p.m.","pm").replace("a.m","am").replace("p.m","pm")
    s=re.sub(r"\s+"," ",s)
    if category=="Time":
        m=re.search(r"(\d{1,2})(?::(\d{2}))?\s*(am|pm)",s)
        if m:
            h=int(m.group(1)); minute=int(m.group(2) or 0); ap=m.group(3)
            if ap=="pm" and h!=12: h+=12
            if ap=="am" and h==12: h=0
            return f"{h:02}:{minute:02}"
    if category=="Price":
        return re.sub(r"[^\d.]","",s)
    if category in ("URL","Email"):
        return s.rstrip("/")
    return re.sub(r"[^a-z0-9]","",s)

def pair_values(src, drf, category):
    results=[]
    remaining=list(drf)

    # First remove equivalent values as PASS.
    unmatched=[]
    for s in src:
        exact=next((d for d in remaining if canon(category,d)==canon(category,s)), None)
        if exact is not None:
            results.append(("PASS",category,s,exact,"Matches approved source."))
            remaining.remove(exact)
        else:
            unmatched.append(s)

    # If source and draft both have a value of a high-risk category, a changed
    # value is a conflict rather than a generic omission.
    for s in unmatched:
        if remaining and category in ("Date","Time","Price"):
            candidate=remaining.pop(0)
            results.append(("CONFLICT",category,s,candidate,"Approved and final values conflict."))
        else:
            results.append(("REVIEW",category,s,None,"Present in approved source but not found in final draft."))

    # Extra high-risk values appearing only in the draft deserve review.
    for d in remaining:
        if category in ("Date","Time","Price"):
            results.append(("REVIEW",category,"Not in approved source",d,"Final draft contains an additional value."))
    return results

def compare(source,draft,required):
    S,D=extract(source),extract(draft)
    rows=[]
    for cat in ["Date","Time","Price","URL","Email"]:
        rows += pair_values(S[cat],D[cat],cat)
    for phrase in required:
        if phrase.lower() in draft.lower():
            rows.append(("PASS","Required wording",phrase,phrase,"Required wording is present."))
        else:
            rows.append(("CONFLICT","Required wording",phrase,"Missing","Required wording is missing."))
    return rows

c1,c2=st.columns(2)
with c1:
    st.subheader("1. Approved source")
    sf=st.file_uploader("Upload approved source",type=["pdf","docx","txt","md"])
    sp=st.text_area("Or paste approved source",height=180)
with c2:
    st.subheader("2. Final draft")
    df=st.file_uploader("Upload final draft",type=["pdf","docx","txt","md"])
    dp=st.text_area("Or paste final draft",height=180)

st.subheader("3. Required wording (optional)")
req=st.text_area("Only add wording that must appear in the final version — one item per line",height=90)

if st.button("Run FinalCheck",type="primary",use_container_width=True):
    source=sp.strip() or extract_text(sf)
    draft=dp.strip() or extract_text(df)
    if not source or not draft:
        st.error("Please provide both documents.")
    else:
        required_items=[x.strip() for x in req.splitlines() if x.strip()]
        source_entities=extract(source)
        draft_entities=extract(draft)
        rows=compare(source,draft,required_items)

        with st.expander("Extraction check — what FinalCheck read from each document"):
            left,right=st.columns(2)
            with left:
                st.markdown("**Approved source**")
                for cat, vals in source_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            with right:
                st.markdown("**Final draft**")
                for cat, vals in draft_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            st.caption("If a fact is missing here, the extraction step—not the comparison step—is the problem.")

        conflicts=[r for r in rows if r[0]=="CONFLICT"]
        reviews=[r for r in rows if r[0]=="REVIEW"]
        passes=[r for r in rows if r[0]=="PASS"]

        a,b,c=st.columns(3)
        a.metric("Conflicts",len(conflicts))
        b.metric("Needs review",len(reviews))
        c.metric("Passed",len(passes))

        if conflicts: st.error("Conflicting information found.")
        elif reviews: st.warning("No direct conflicts found, but some source information needs review.")
        else: st.success("No factual conflicts detected.")

        for status,cat,s,d,msg in conflicts+reviews+passes:
            icon={"CONFLICT":"🔴","REVIEW":"🟡","PASS":"🟢"}[status]
            with st.expander(f"{icon} {status} · {cat} · {s}", expanded=status=="CONFLICT"):
                st.write(f"**Approved source:** {s}")
                st.write(f"**Final draft:** {d or 'Not found'}")
                st.write(msg)

        report=[{"status":x[0],"category":x[1],"approved":x[2],"draft":x[3],"reason":x[4]} for x in rows]
        st.download_button("Download QA report",json.dumps(report,indent=2,ensure_ascii=False),
                           "finalcheck_v2_report.json","application/json",use_container_width=True)

st.divider()
st.caption("V2 separates direct conflicts from omissions/review items and recognizes equivalent time formatting.")
import streamlit as st
import re, json
from io import BytesIO
from difflib import SequenceMatcher

st.set_page_config(page_title="FinalCheck V2", page_icon="✓", layout="wide")

st.markdown("""
<style>
.block-container {max-width: 1120px; padding-top: 2rem;}
.pass {padding:14px;border-radius:10px;background:#eaf7ee;margin:8px 0}
.conflict {padding:14px;border-radius:10px;background:#fdecec;margin:8px 0}
.review {padding:14px;border-radius:10px;background:#fff6df;margin:8px 0}
</style>
""", unsafe_allow_html=True)

st.title("FinalCheck")
st.caption("Pre-flight fact checking for communications documents — V2")

def extract_text(uploaded):
    if uploaded is None: return ""
    data = uploaded.getvalue()
    name = uploaded.name.lower()
    if name.endswith((".txt",".md")):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        from docx import Document
        return "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        return "\n".join((p.extract_text() or "") for p in PdfReader(BytesIO(data)).pages)
    return ""

MONTHS = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"

def unique(xs):
    seen=[]; keys=set()
    for x in xs:
        x=re.sub(r"\s+"," ",x.strip().rstrip(".,;:"))
        k=x.lower()
        if x and k not in keys:
            keys.add(k); seen.append(x)
    return seen

def extract(text):
    return {
      "Date": unique(re.findall(rf"\b{MONTHS}\s+\d{{1,2}}(?:,\s*\d{{4}})?\b|\b\d{{4}}[-/]\d{{1,2}}[-/]\d{{1,2}}\b", text, re.I)),
      "Time": unique(re.findall(r"\b\d{1,2}(?::\d{2})?\s*(?:a\.?m\.?|p\.?m\.?)\b", text, re.I)),
      "Price": unique(re.findall(r"(?:CAD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?", text, re.I)),
      "URL": unique(re.findall(r"https?://[^\s<>\]\)]+|www\.[^\s<>\]\)]+", text, re.I)),
      "Email": unique(re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)),
    }

def canon(category, value):
    s=value.lower().strip()
    s=s.replace("a.m.","am").replace("p.m.","pm").replace("a.m","am").replace("p.m","pm")
    s=re.sub(r"\s+"," ",s)
    if category=="Time":
        m=re.search(r"(\d{1,2})(?::(\d{2}))?\s*(am|pm)",s)
        if m:
            h=int(m.group(1)); minute=int(m.group(2) or 0); ap=m.group(3)
            if ap=="pm" and h!=12: h+=12
            if ap=="am" and h==12: h=0
            return f"{h:02}:{minute:02}"
    if category=="Price":
        return re.sub(r"[^\d.]","",s)
    if category in ("URL","Email"):
        return s.rstrip("/")
    return re.sub(r"[^a-z0-9]","",s)

def pair_values(src, drf, category):
    results=[]
    remaining=list(drf)

    # First remove equivalent values as PASS.
    unmatched=[]
    for s in src:
        exact=next((d for d in remaining if canon(category,d)==canon(category,s)), None)
        if exact is not None:
            results.append(("PASS",category,s,exact,"Matches approved source."))
            remaining.remove(exact)
        else:
            unmatched.append(s)

    # If source and draft both have a value of a high-risk category, a changed
    # value is a conflict rather than a generic omission.
    for s in unmatched:
        if remaining and category in ("Date","Time","Price"):
            candidate=remaining.pop(0)
            results.append(("CONFLICT",category,s,candidate,"Approved and final values conflict."))
        else:
            results.append(("REVIEW",category,s,None,"Present in approved source but not found in final draft."))

    # Extra high-risk values appearing only in the draft deserve review.
    for d in remaining:
        if category in ("Date","Time","Price"):
            results.append(("REVIEW",category,"Not in approved source",d,"Final draft contains an additional value."))
    return results

def compare(source,draft,required):
    S,D=extract(source),extract(draft)
    rows=[]
    for cat in ["Date","Time","Price","URL","Email"]:
        rows += pair_values(S[cat],D[cat],cat)
    for phrase in required:
        if phrase.lower() in draft.lower():
            rows.append(("PASS","Required wording",phrase,phrase,"Required wording is present."))
        else:
            rows.append(("CONFLICT","Required wording",phrase,"Missing","Required wording is missing."))
    return rows

c1,c2=st.columns(2)
with c1:
    st.subheader("1. Approved source")
    sf=st.file_uploader("Upload approved source",type=["pdf","docx","txt","md"])
    sp=st.text_area("Or paste approved source",height=180)
with c2:
    st.subheader("2. Final draft")
    df=st.file_uploader("Upload final draft",type=["pdf","docx","txt","md"])
    dp=st.text_area("Or paste final draft",height=180)

st.subheader("3. Required wording (optional)")
req=st.text_area("Only add wording that must appear in the final version — one item per line",height=90)

if st.button("Run FinalCheck",type="primary",use_container_width=True):
    source=sp.strip() or extract_text(sf)
    draft=dp.strip() or extract_text(df)
    if not source or not draft:
        st.error("Please provide both documents.")
    else:
        required_items=[x.strip() for x in req.splitlines() if x.strip()]
        source_entities=extract(source)
        draft_entities=extract(draft)
        rows=compare(source,draft,required_items)

        with st.expander("Extraction check — what FinalCheck read from each document"):
            left,right=st.columns(2)
            with left:
                st.markdown("**Approved source**")
                for cat, vals in source_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            with right:
                st.markdown("**Final draft**")
                for cat, vals in draft_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            st.caption("If a fact is missing here, the extraction step—not the comparison step—is the problem.")

        conflicts=[r for r in rows if r[0]=="CONFLICT"]
        reviews=[r for r in rows if r[0]=="REVIEW"]
        passes=[r for r in rows if r[0]=="PASS"]

        a,b,c=st.columns(3)
        a.metric("Conflicts",len(conflicts))
        b.metric("Needs review",len(reviews))
        c.metric("Passed",len(passes))

        if conflicts: st.error("Conflicting information found.")
        elif reviews: st.warning("No direct conflicts found, but some source information needs review.")
        else: st.success("No factual conflicts detected.")

        for status,cat,s,d,msg in conflicts+reviews+passes:
            icon={"CONFLICT":"🔴","REVIEW":"🟡","PASS":"🟢"}[status]
            with st.expander(f"{icon} {status} · {cat} · {s}", expanded=status=="CONFLICT"):
                st.write(f"**Approved source:** {s}")
                st.write(f"**Final draft:** {d or 'Not found'}")
                st.write(msg)

        report=[{"status":x[0],"category":x[1],"approved":x[2],"draft":x[3],"reason":x[4]} for x in rows]
        st.download_button("Download QA report",json.dumps(report,indent=2,ensure_ascii=False),
                           "finalcheck_v2_report.json","application/json",use_container_width=True)

st.divider()
st.caption("V2 separates direct conflicts from omissions/review items and recognizes equivalent time formatting.")
import streamlit as st
import re, json
from io import BytesIO
from difflib import SequenceMatcher

st.set_page_config(page_title="FinalCheck V2", page_icon="✓", layout="wide")

st.markdown("""
<style>
.block-container {max-width: 1120px; padding-top: 2rem;}
.pass {padding:14px;border-radius:10px;background:#eaf7ee;margin:8px 0}
.conflict {padding:14px;border-radius:10px;background:#fdecec;margin:8px 0}
.review {padding:14px;border-radius:10px;background:#fff6df;margin:8px 0}
</style>
""", unsafe_allow_html=True)

st.title("FinalCheck")
st.caption("Pre-flight fact checking for communications documents — V2")

def extract_text(uploaded):
    if uploaded is None: return ""
    data = uploaded.getvalue()
    name = uploaded.name.lower()
    if name.endswith((".txt",".md")):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        from docx import Document
        return "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        return "\n".join((p.extract_text() or "") for p in PdfReader(BytesIO(data)).pages)
    return ""

MONTHS = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"

def unique(xs):
    seen=[]; keys=set()
    for x in xs:
        x=re.sub(r"\s+"," ",x.strip().rstrip(".,;:"))
        k=x.lower()
        if x and k not in keys:
            keys.add(k); seen.append(x)
    return seen

def extract(text):
    return {
      "Date": unique(re.findall(rf"\b{MONTHS}\s+\d{{1,2}}(?:,\s*\d{{4}})?\b|\b\d{{4}}[-/]\d{{1,2}}[-/]\d{{1,2}}\b", text, re.I)),
      "Time": unique(re.findall(r"\b\d{1,2}(?::\d{2})?\s*(?:a\.?m\.?|p\.?m\.?)\b", text, re.I)),
      "Price": unique(re.findall(r"(?:CAD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?", text, re.I)),
      "URL": unique(re.findall(r"https?://[^\s<>\]\)]+|www\.[^\s<>\]\)]+", text, re.I)),
      "Email": unique(re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)),
    }

def canon(category, value):
    s=value.lower().strip()
    s=s.replace("a.m.","am").replace("p.m.","pm").replace("a.m","am").replace("p.m","pm")
    s=re.sub(r"\s+"," ",s)
    if category=="Time":
        m=re.search(r"(\d{1,2})(?::(\d{2}))?\s*(am|pm)",s)
        if m:
            h=int(m.group(1)); minute=int(m.group(2) or 0); ap=m.group(3)
            if ap=="pm" and h!=12: h+=12
            if ap=="am" and h==12: h=0
            return f"{h:02}:{minute:02}"
    if category=="Price":
        return re.sub(r"[^\d.]","",s)
    if category in ("URL","Email"):
        return s.rstrip("/")
    return re.sub(r"[^a-z0-9]","",s)

def pair_values(src, drf, category):
    results=[]
    remaining=list(drf)

    # First remove equivalent values as PASS.
    unmatched=[]
    for s in src:
        exact=next((d for d in remaining if canon(category,d)==canon(category,s)), None)
        if exact is not None:
            results.append(("PASS",category,s,exact,"Matches approved source."))
            remaining.remove(exact)
        else:
            unmatched.append(s)

    # If source and draft both have a value of a high-risk category, a changed
    # value is a conflict rather than a generic omission.
    for s in unmatched:
        if remaining and category in ("Date","Time","Price"):
            candidate=remaining.pop(0)
            results.append(("CONFLICT",category,s,candidate,"Approved and final values conflict."))
        else:
            results.append(("REVIEW",category,s,None,"Present in approved source but not found in final draft."))

    # Extra high-risk values appearing only in the draft deserve review.
    for d in remaining:
        if category in ("Date","Time","Price"):
            results.append(("REVIEW",category,"Not in approved source",d,"Final draft contains an additional value."))
    return results

def compare(source,draft,required):
    S,D=extract(source),extract(draft)
    rows=[]
    for cat in ["Date","Time","Price","URL","Email"]:
        rows += pair_values(S[cat],D[cat],cat)
    for phrase in required:
        if phrase.lower() in draft.lower():
            rows.append(("PASS","Required wording",phrase,phrase,"Required wording is present."))
        else:
            rows.append(("CONFLICT","Required wording",phrase,"Missing","Required wording is missing."))
    return rows

c1,c2=st.columns(2)
with c1:
    st.subheader("1. Approved source")
    sf=st.file_uploader("Upload approved source",type=["pdf","docx","txt","md"])
    sp=st.text_area("Or paste approved source",height=180)
with c2:
    st.subheader("2. Final draft")
    df=st.file_uploader("Upload final draft",type=["pdf","docx","txt","md"])
    dp=st.text_area("Or paste final draft",height=180)

st.subheader("3. Required wording (optional)")
req=st.text_area("Only add wording that must appear in the final version — one item per line",height=90)

if st.button("Run FinalCheck",type="primary",use_container_width=True):
    source=sp.strip() or extract_text(sf)
    draft=dp.strip() or extract_text(df)
    if not source or not draft:
        st.error("Please provide both documents.")
    else:
        required_items=[x.strip() for x in req.splitlines() if x.strip()]
        source_entities=extract(source)
        draft_entities=extract(draft)
        rows=compare(source,draft,required_items)

        with st.expander("Extraction check — what FinalCheck read from each document"):
            left,right=st.columns(2)
            with left:
                st.markdown("**Approved source**")
                for cat, vals in source_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            with right:
                st.markdown("**Final draft**")
                for cat, vals in draft_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            st.caption("If a fact is missing here, the extraction step—not the comparison step—is the problem.")

        conflicts=[r for r in rows if r[0]=="CONFLICT"]
        reviews=[r for r in rows if r[0]=="REVIEW"]
        passes=[r for r in rows if r[0]=="PASS"]

        a,b,c=st.columns(3)
        a.metric("Conflicts",len(conflicts))
        b.metric("Needs review",len(reviews))
        c.metric("Passed",len(passes))

        if conflicts: st.error("Conflicting information found.")
        elif reviews: st.warning("No direct conflicts found, but some source information needs review.")
        else: st.success("No factual conflicts detected.")

        for status,cat,s,d,msg in conflicts+reviews+passes:
            icon={"CONFLICT":"🔴","REVIEW":"🟡","PASS":"🟢"}[status]
            with st.expander(f"{icon} {status} · {cat} · {s}", expanded=status=="CONFLICT"):
                st.write(f"**Approved source:** {s}")
                st.write(f"**Final draft:** {d or 'Not found'}")
                st.write(msg)

        report=[{"status":x[0],"category":x[1],"approved":x[2],"draft":x[3],"reason":x[4]} for x in rows]
        st.download_button("Download QA report",json.dumps(report,indent=2,ensure_ascii=False),
                           "finalcheck_v2_report.json","application/json",use_container_width=True)

st.divider()
st.caption("V2 separates direct conflicts from omissions/review items and recognizes equivalent time formatting.")
import streamlit as st
import re, json
from io import BytesIO
from difflib import SequenceMatcher

st.set_page_config(page_title="FinalCheck V2", page_icon="✓", layout="wide")

st.markdown("""
<style>
.block-container {max-width: 1120px; padding-top: 2rem;}
.pass {padding:14px;border-radius:10px;background:#eaf7ee;margin:8px 0}
.conflict {padding:14px;border-radius:10px;background:#fdecec;margin:8px 0}
.review {padding:14px;border-radius:10px;background:#fff6df;margin:8px 0}
</style>
""", unsafe_allow_html=True)

st.title("FinalCheck")
st.caption("Pre-flight fact checking for communications documents — V2")

def extract_text(uploaded):
    if uploaded is None: return ""
    data = uploaded.getvalue()
    name = uploaded.name.lower()
    if name.endswith((".txt",".md")):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".docx"):
        from docx import Document
        return "\n".join(p.text for p in Document(BytesIO(data)).paragraphs)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        return "\n".join((p.extract_text() or "") for p in PdfReader(BytesIO(data)).pages)
    return ""

MONTHS = r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"

def unique(xs):
    seen=[]; keys=set()
    for x in xs:
        x=re.sub(r"\s+"," ",x.strip().rstrip(".,;:"))
        k=x.lower()
        if x and k not in keys:
            keys.add(k); seen.append(x)
    return seen

def extract(text):
    return {
      "Date": unique(re.findall(rf"\b{MONTHS}\s+\d{{1,2}}(?:,\s*\d{{4}})?\b|\b\d{{4}}[-/]\d{{1,2}}[-/]\d{{1,2}}\b", text, re.I)),
      "Time": unique(re.findall(r"\b\d{1,2}(?::\d{2})?\s*(?:a\.?m\.?|p\.?m\.?)\b", text, re.I)),
      "Price": unique(re.findall(r"(?:CAD\s*)?\$\s?\d[\d,]*(?:\.\d{1,2})?", text, re.I)),
      "URL": unique(re.findall(r"https?://[^\s<>\]\)]+|www\.[^\s<>\]\)]+", text, re.I)),
      "Email": unique(re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)),
    }

def canon(category, value):
    s=value.lower().strip()
    s=s.replace("a.m.","am").replace("p.m.","pm").replace("a.m","am").replace("p.m","pm")
    s=re.sub(r"\s+"," ",s)
    if category=="Time":
        m=re.search(r"(\d{1,2})(?::(\d{2}))?\s*(am|pm)",s)
        if m:
            h=int(m.group(1)); minute=int(m.group(2) or 0); ap=m.group(3)
            if ap=="pm" and h!=12: h+=12
            if ap=="am" and h==12: h=0
            return f"{h:02}:{minute:02}"
    if category=="Price":
        return re.sub(r"[^\d.]","",s)
    if category in ("URL","Email"):
        return s.rstrip("/")
    return re.sub(r"[^a-z0-9]","",s)

def pair_values(src, drf, category):
    results=[]
    remaining=list(drf)

    # First remove equivalent values as PASS.
    unmatched=[]
    for s in src:
        exact=next((d for d in remaining if canon(category,d)==canon(category,s)), None)
        if exact is not None:
            results.append(("PASS",category,s,exact,"Matches approved source."))
            remaining.remove(exact)
        else:
            unmatched.append(s)

    # If source and draft both have a value of a high-risk category, a changed
    # value is a conflict rather than a generic omission.
    for s in unmatched:
        if remaining and category in ("Date","Time","Price"):
            candidate=remaining.pop(0)
            results.append(("CONFLICT",category,s,candidate,"Approved and final values conflict."))
        else:
            results.append(("REVIEW",category,s,None,"Present in approved source but not found in final draft."))

    # Extra high-risk values appearing only in the draft deserve review.
    for d in remaining:
        if category in ("Date","Time","Price"):
            results.append(("REVIEW",category,"Not in approved source",d,"Final draft contains an additional value."))
    return results

def compare(source,draft,required):
    S,D=extract(source),extract(draft)
    rows=[]
    for cat in ["Date","Time","Price","URL","Email"]:
        rows += pair_values(S[cat],D[cat],cat)
    for phrase in required:
        if phrase.lower() in draft.lower():
            rows.append(("PASS","Required wording",phrase,phrase,"Required wording is present."))
        else:
            rows.append(("CONFLICT","Required wording",phrase,"Missing","Required wording is missing."))
    return rows

c1,c2=st.columns(2)
with c1:
    st.subheader("1. Approved source")
    sf=st.file_uploader("Upload approved source",type=["pdf","docx","txt","md"])
    sp=st.text_area("Or paste approved source",height=180)
with c2:
    st.subheader("2. Final draft")
    df=st.file_uploader("Upload final draft",type=["pdf","docx","txt","md"])
    dp=st.text_area("Or paste final draft",height=180)

st.subheader("3. Required wording (optional)")
req=st.text_area("Only add wording that must appear in the final version — one item per line",height=90)

if st.button("Run FinalCheck",type="primary",use_container_width=True):
    source=sp.strip() or extract_text(sf)
    draft=dp.strip() or extract_text(df)
    if not source or not draft:
        st.error("Please provide both documents.")
    else:
        required_items=[x.strip() for x in req.splitlines() if x.strip()]
        source_entities=extract(source)
        draft_entities=extract(draft)
        rows=compare(source,draft,required_items)

        with st.expander("Extraction check — what FinalCheck read from each document"):
            left,right=st.columns(2)
            with left:
                st.markdown("**Approved source**")
                for cat, vals in source_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            with right:
                st.markdown("**Final draft**")
                for cat, vals in draft_entities.items():
                    st.write(f"{cat}: " + (", ".join(vals) if vals else "—"))
            st.caption("If a fact is missing here, the extraction step—not the comparison step—is the problem.")

        conflicts=[r for r in rows if r[0]=="CONFLICT"]
        reviews=[r for r in rows if r[0]=="REVIEW"]
        passes=[r for r in rows if r[0]=="PASS"]

        a,b,c=st.columns(3)
        a.metric("Conflicts",len(conflicts))
        b.metric("Needs review",len(reviews))
        c.metric("Passed",len(passes))

        if conflicts: st.error("Conflicting information found.")
        elif reviews: st.warning("No direct conflicts found, but some source information needs review.")
        else: st.success("No factual conflicts detected.")

        for status,cat,s,d,msg in conflicts+reviews+passes:
            icon={"CONFLICT":"🔴","REVIEW":"🟡","PASS":"🟢"}[status]
            with st.expander(f"{icon} {status} · {cat} · {s}", expanded=status=="CONFLICT"):
                st.write(f"**Approved source:** {s}")
                st.write(f"**Final draft:** {d or 'Not found'}")
                st.write(msg)

        report=[{"status":x[0],"category":x[1],"approved":x[2],"draft":x[3],"reason":x[4]} for x in rows]
        st.download_button("Download QA report",json.dumps(report,indent=2,ensure_ascii=False),
                           "finalcheck_v2_report.json","application/json",use_container_width=True)

st.divider()
st.caption("V2 separates direct conflicts from omissions/review items and recognizes equivalent time formatting.")
